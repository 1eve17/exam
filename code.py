"""
一.注册与登陆：
注册说明：
1.一开始的用户信息是空的，用户信息文件由学校给出，由学校指派的超级管理员提前录入到数据库中去，录入一一个xlsx格式的表格，格式为：第一列为学号/工号，第二列为用户种类 student / teacher / admins / nt_authority
2.此后，在注册时需要输入学号/工号，用户名与密码，系统会自动判断注册者是否为本学校/组织的人，注册完成后，以后登陆时只需要输入用户名与密码即可
3.考虑到安全性问题，用户信息仅可导入，不可导出，需要更新用户信息就需要将原本的数据库删除之后再重新导入
二.题目导入规则：
如果是docx或者txt文件
1.必须以 #单选题# / #多选题# / #判断题# / #填空题# / #主观题# 来开头，这些字符会被吸收掉，然后就开始录入题干
2.每条题目以一行回车分割，答案在最后一行
如果是excel：
1.第一列是题干，必须以 #单选题# / #多选题# / #判断题# / #填空题# / #主观题# 来开头，这些字符会被吸收掉，然后就开始录入题干
2.如果是主观题，第二列是答案
3.如果是选择题，第2-7行是选项，
一开始题库是空的在data里面有一些示例题目，学校的管理员可以将它们拖动到同层文件夹下导入。
三.权限分级：
游客模式:只能查看题目（几万位）
student:只能查看题目（几千位）
teacher:可以查看以及修改题目（几百位）
admin:可以导入导出题目（几十位）
nt_authority:可以修改用户的信息（几位）
四.请不要随意改变后端文件的名称，否则会导致出错
                                                                        ——by Level7
"""
import time  # 在各种回显时给予用户反应时间
import tkinter  # 主要的GUI
from PIL import Image, ImageTk
from tkinter import scrolledtext  # 我也不知道为什么这个组件识别不到，需要单独引用
import sqlite3  # 数据库
import docx  # 读取.docx文档
import docx.opc.exceptions  # 一样,需要单独引用,报错检测用
from docx import Document  # 导出到docx时写入用
import openpyxl  # 读取.xlsx文件
import sys  # 为了封装起来是exit按钮还有效
from tkinter import ttk


class Color:
    # 不是常见的rgb调色法，不过也差不多,去用PS的调色板可以快速调出更多这里的颜色及对应值
    white = '#FFFFFF'  # 白
    black = '#000000'  # 黑
    red = '#FF0000'  # 红
    green = '#00FF00'  # 绿
    blue = '#0000FF'  # 蓝
    yellow = '#FFFF00'  # 黄
    purple = '#FF00FF'  # 紫
    grey = '#DBDAD5'  # 灰
    sky_blue = '#00FFFF'  # 天空蓝
    sky_blue_out = '#11809c'  # 淡天空蓝
    b_c1 = '#9bd5eb'  # 背景色1 中
    b_c2 = '#9394ee'  # 背景色2 深
    b_c3 = '#e7e6f6'  # 背景色3 淡
    b_c4 = '#dbd9ee'  # 背景色4 偏黑的b_c3


def renovate():
    """
    每次进入新的界面时先刷新一下，重新摆放所有最基本的元素，因此在所有窗口转换的函数的第一步就是调用此函数
    对于其他的控件，要由其他函数自己来解决
    :return: None
    """
    global screen
    # 使用画布来遮住所有控件
    tkinter.Canvas(screen, height=714, width=1344, bg='white').place(x=0, y=0)
    # 重新摆放背景图片
    tkinter.Label(screen, image=background_image, bg=Color.b_c3).place(x=0, y=0)
    # 摆放一个home按钮，使之链接到主菜单
    tkinter.Button(screen, text="Home", font=("SimSun", 30), command=lambda: main_menu(authority), bg=Color.sky_blue,
                   activebackground=Color.sky_blue_out).place(x=1230, y=5)
    return screen  # 记得要用screen来接收返回值，即screen = renovate()


def read_file(file_num):
    """
    读取数据库对应的表
    :param file_num: 对应的文件编号,见下面的字典dict_name
    :return: 形如[[id(int),题干(str),选项[list(str)],答案(str)],...]的三维列表结构
    """
    connection = sqlite3.connect('data/examination_questions.db')  # 开启数据库的链接
    cur = connection.cursor()  # 创建游标
    # 创建字典，用格式化的字符串与之结合减少代码的重复量
    dict_name = {0: "radio", 1: "multiple_choice", 2: "true_or_false", 3: "fill_in_the_blank", 4: "subjective", }
    # 下面的try是为了反之用户在数据库为空的时候访问造成的报错
    try:
        list_q = cur.execute(f"SELECT * FROM {dict_name[file_num]}").fetchall()  # 将读取到的问题存放到list_q里面
    # 如果题库为空，那么就提醒用户，然后在1S后返回主菜单
    except sqlite3.OperationalError:
        list_q = None
        tkinter.Label(screen, text="题库中还没有这种题!", font=("SimSun", 50), bg=Color.b_c3).place(x=300, y=250)
        screen.update()
        time.sleep(1)
        main_menu(authority)
    cur.close()  # 关闭游标
    connection.close()  # 关闭数据库
    # 下面对list_q规整化，使之变成形如[[id(int),题干(str),选项[list(str)],答案(str)],...]的三维列表结构
    list_retuning = []
    for i in list_q:
        list_temp = []
        list_temp.append(i[0])  # id
        list_temp.append(i[1])  # 题干
        if file_num < 2:  # 选项
            list_temp.append(eval(i[2]))
        list_temp.append(i[-1])  # 答案
        list_retuning.append(list_temp)
    return list_retuning


def rewrite_file(file_num, types, list_content=None, id = None):
    """
    对数据库的增删改的统一函数
    :param file_num: 操作的文件对象
    :param types: 操作类型，只能为字符串add,sub或revise，将其他的参数视为严重错误，以后维护代码时也不该忘记这一点
    :param list_content: 要修改/增加的内容
    :param id: 题目的id
    :return: None
    """
    # 创建字典，用格式化的字符串与之结合减少代码的重复量
    dict_name = {0: "radio", 1: "multiple_choice", 2: "true_or_false", 3: "fill_in_the_blank", 4: "subjective"}
    connection = sqlite3.connect(f'data/examination_questions.db')  # 链接数据库
    cur = connection.cursor()  # 创建游标
    # 用户选择了添加题目的操作
    if types == "add":
        # 创造一个新的且绝对不重复的id,是当前最大id+1,这个算法可以改进为id不废弃的类型，这里节约算力，减少代码量，所有没有这么做
        id = max([i[0] for i in cur.execute(f"SELECT id FROM {dict_name[file_num]}").fetchall()]) + 1
        if file_num < 2:  # 添加选择题
            cur.execute(f"INSERT INTO {dict_name[file_num]} values(?,?,?,?)",
                        (id, list_content[0], str(list_content[1]), list_content[-1]))
        else:  # 添加主观题
            cur.execute(f"INSERT INTO {dict_name[file_num]} values(?,?,?)",
                        (id, list_content[0], list_content[-1]))
        # 数据库结束三连:确认修改,关闭游标,关闭链接
        connection.commit()
        cur.close()
        connection.close()
        add(file_num)
    # 用户选择了修改题目的操作，我们的思想是无论删除还是修改，我们都删除原来的题目，如果是修改的话我们再将用户的提交的新的信息填到原来删除的空位里
    elif types == "sub" or types == "revise":
        cur.execute(f"DELETE FROM {dict_name[file_num]} WHERE id={id}")  # 直接删除题目
        # 对于修改题目，
        if types == "revise":
            if file_num < 2:  # 修改选择题
                cur.execute(f"INSERT INTO {dict_name[file_num]} values(?,?,?,?)",
                            (id, list_content[0], str(list_content[1]), list_content[-1]))
            else:  # 修改主观题
                cur.execute(f"INSERT INTO {dict_name[file_num]} values(?,?,?)",
                            (id, list_content[0], list_content[-1]))
        # 数据库结束三连:确认修改,关闭游标,关闭链接
        connection.commit()
        cur.close()
        connection.close()
        sub_and_revise(file_num)
    else:
        # 将参数错误视为严重错误
        print("错误参数types，于rewrite_file中。")
        exit()


def import_member_information(file_path=None):
    screen = renovate()
    tkinter.Button(screen, text="←", font=("SimSun", 30), command=lambda: main_menu(authority), bg=Color.sky_blue,
                   activebackground=Color.sky_blue_out).place(x=5, y=5)

    # 初次进入这个窗口
    if file_path is None:
        # 提示字段
        tkinter.Label(screen, text="请输入要导入的用户信息文件名,请确保在data文件夹中且必须为xlsx格式!", font=font1,
                      bg=Color.b_c3).place(x=400, y=200)
        # 输入框与确定按钮
        file_name_text = tkinter.Entry(screen, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        file_name_text.place(x=400, y=244)
        tkinter.Button(screen, text="确认导入", font=font1,
                       command=lambda: import_member_information(file_name_text.get()),
                       bg=Color.b_c1, activebackground=Color.b_c2).place(x=400, y=280)
    else:
        list_org_info = []
        try:
            file1 = openpyxl.load_workbook(f"data/{file_path}").active
        except Exception as e:
            file1 = None
            # tkinter.Label(screen, text="输入有误!", font=("SimSun", 50), bg=Color.b_c3).place(x=300, y=250)
            tkinter.Label(screen, text=f"文件错误: {str(e)}", font=("SimSun", 20), bg=Color.b_c3).place(x=150, y=250)
            screen.update()
            time.sleep(1)
            import_member_information()
        for line in file1:
            list_temp = []
            for val in line:
                list_temp.append(val.value)
            list_org_info.append(list_temp)
        connection = sqlite3.connect('data/usersinfo.db')  # 在同级目录下如果没有名为 usersinfo.db 的数据库，那么就创建一个
        cur = connection.cursor()  # 创建游标
        cur.execute("CREATE TABLE IF NOT EXISTS org_info(id TEXT PRIMARY KEY, types TEXT)")  # 创建用户信息表,包含id，用户名，密码三个数据类型
        try:
            cur.executemany('INSERT INTO org_info VALUES (?,?)', list_org_info)
        except sqlite3.IntegrityError:
            pass
            tkinter.Label(screen, text="存在学号/工号重复的情况!", font=("SimSun", 50), bg=Color.b_c3).place(x=300,
                                                                                                             y=250)
            screen.update()
            time.sleep(1)
            main_menu(authority)

        # 结束五连:确认修改,关闭游标,关闭链接，提示用户，返回主菜单
        connection.commit()
        cur.close()
        connection.close()
        tkinter.Label(screen, text="导入成功!", font=("SimSun", 50), bg=Color.b_c3).place(x=300, y=250)
        screen.update()
        time.sleep(1)
        main_menu(authority)

def search_questions_by_keyword(keyword):
    """
    在题库中按关键词查找相关题目和答案
    :param keyword: str 用户输入的关键词
    :return: None, 在界面显示结果
    """
    screen = renovate()
    tkinter.Button(screen, text="←", font=("SimSun", 30), command=lambda: main_menu(authority), bg=Color.sky_blue,
                   activebackground=Color.sky_blue_out).place(x=5, y=5)

    if not keyword:
        tkinter.Label(screen, text="请输入关键词！", font=font1, bg=Color.b_c3).place(x=400, y=200)
        screen.update()
        time.sleep(1)
        main_menu(authority)
        return

    # 连接数据库，查找包含关键词的题目或答案
    connection = sqlite3.connect('data/examination_questions.db')
    cur = connection.cursor()

    dict_name = {0: "radio", 1: "multiple_choice", 2: "true_or_false", 3: "fill_in_the_blank", 4: "subjective"}
    results = []

    for key, table_name in dict_name.items():
        try:
            # 查找题干和答案中包含关键词的题目
            rows = cur.execute(f"SELECT * FROM {table_name} WHERE TRIM(problem) LIKE ? OR TRIM(answer) LIKE ?",
                               (f"%{keyword}%", f"%{keyword}%")).fetchall()
            for row in rows:
                results.append((table_name, row))
        except sqlite3.OperationalError:
            continue

    cur.close()
    connection.close()

    # 显示搜索结果
    if not results:
        tkinter.Label(screen, text="未找到相关题目！", font=font1, bg=Color.b_c3).place(x=400, y=200)
        screen.update()
        time.sleep(1)
        main_menu(authority)
        return

    tkinter.Label(screen, text=f"找到 {len(results)} 个相关题目：", font=font1, bg=Color.b_c3).place(x=400, y=50)

    # 创建一个Canvas容器并添加滚动条
    canvas = tkinter.Canvas(screen, height=500, width=1200)
    canvas.place(x=50, y=100)

    scrollbar = tkinter.Scrollbar(screen, orient="vertical", command=canvas.yview)
    scrollbar.place(x=1250, y=100, height=500)
    canvas.config(yscrollcommand=scrollbar.set)

    # 创建一个Frame来容纳所有搜索结果标签
    result_frame = tkinter.Frame(canvas)

    # 将结果放入到frame中
    for i, (table_name, row) in enumerate(results):
        problem = row[1]  # 题干
        options = row[2] if table_name in ["radio", "multiple_choice"] else ""
        answer = row[-1]  # 答案

        display_text = f"[{table_name}] ID: {row[0]} 题干: {problem} 答案: {answer}"
        if options:
            display_text += f" 选项: {options}"

        # 创建标签并将其添加到Frame中
        result_label = tkinter.Label(result_frame, text=display_text, font=("SimSun", 12), bg=Color.b_c3, wraplength=1000, justify="left")
        result_label.grid(row=i, column=0, sticky="w", pady=5)

    # 将frame添加到canvas中并设置滚动区域
    canvas.create_window((0, 0), window=result_frame, anchor="nw")

    # 更新滚动区域的大小
    result_frame.update_idletasks()
    canvas.config(scrollregion=canvas.bbox("all"))

    # 绑定鼠标滚轮事件，使其可以滚动Canvas内容
    def on_mouse_wheel(event):
        canvas.yview_scroll(int(-1*(event.delta/120)), "units")

    canvas.bind_all("<MouseWheel>", on_mouse_wheel)  # 绑定鼠标滚轮事件

    screen.update()


def open_keyword_search():
    """
    打开关键词输入界面
    """
    screen = renovate()
    tkinter.Button(screen, text="←", font=("SimSun", 30), command=lambda: main_menu(authority), bg=Color.sky_blue,
                   activebackground=Color.sky_blue_out).place(x=5, y=5)

    tkinter.Label(screen, text="请输入关键词：", font=font1, bg=Color.b_c3).place(x=400, y=200)

    keyword_entry = tkinter.Entry(screen, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
    keyword_entry.place(x=400, y=250)

    tkinter.Button(screen, text="查找", font=font1, command=lambda: search_questions_by_keyword(keyword_entry.get()),
                   bg=Color.sky_blue, activebackground=Color.sky_blue_out).place(x=400, y=300)

def register(userid=None, username=None, password=None):
    """
    注册界面，
    :return: None，但是记得注册完之后返回login界面
    """
    # 背景
    tkinter.Label(screen, image=background_image).place(x=0, y=0)  # 设置背景
    # login_background_image = tkinter.PhotoImage(file='login_background_image.png')
    img = Image.open("data\login_background_image.png")  # 使用pillow打开PNG图片
    login_background_image = ImageTk.PhotoImage(img)
    # 创建标签并设置背景图
    label_login_background_image = tkinter.Label(screen, image=login_background_image)
    label_login_background_image.place(x=383, y=204)
    # 初次进入该界面
    if userid is None and password is None and username is None:
        # 注册界面元素
        tkinter.Label(screen, text="注册", font=font1, bg=Color.b_c3).place(x=650, y=170)

        tkinter.Label(screen, text="学号/工号:", font=font1, bg=Color.b_c3).place(x=415, y=220)
        entry_userid = tkinter.Entry(screen, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        entry_userid.place(x=555, y=220)

        tkinter.Label(screen, text="   用户名:", font=font1, bg=Color.b_c3).place(x=415, y=270)
        entry_username = tkinter.Entry(screen, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        entry_username.place(x=555, y=270)

        tkinter.Label(screen, text="     密码:", font=font1, bg=Color.b_c3).place(x=415, y=320)
        entry_password = tkinter.Entry(screen, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        entry_password.place(x=555, y=320)

        # 注册按钮
        register_button = tkinter.Button(screen, text="注册", font=font1, activebackground=Color.sky_blue,
                                         bg=Color.white,
                                         command=lambda: register(entry_userid.get(), entry_username.get(),
                                                                  entry_password.get()))
        register_button.place(x=720, y=374)

        # 返回登录界面按钮，差点忘了这个
        back_to_login_button = tkinter.Button(screen, text="返回登录", font=font1, activebackground=Color.sky_blue,
                                              bg=Color.white,
                                              command=lambda: login(None, None, None))
        back_to_login_button.place(x=550, y=374)

        # 这里也要循环显示才行
        screen.mainloop()
    # 用户已经输入id与密码
    else:
        connection = sqlite3.connect('data/usersinfo.db')  # 在同级目录下如果没有名为 usersinfo.db 的数据库，那么就创建一个
        cur = connection.cursor()  # 创建游标
        # 判断用户id是否在学校/组织给出的文件里面
        yes_or_no = False
        user_types = None
        cur.execute("CREATE TABLE IF NOT EXISTS org_info(id TEXT PRIMARY KEY, types TEXT)")  # 防止报错
        list_org_info = cur.execute("SELECT * FROM org_info").fetchall()
        for i in list_org_info:
            if i[0] == str(userid):
                user_types = i[1]
                yes_or_no = True
        # 如果不在学校的文件里面，就提醒用户使用游客模式访问
        if not yes_or_no:
            connection.commit()
            cur.close()
            connection.close()
            tkinter.Label(screen, text="您不属于系统用户，请使用游客模式访问!", font=("SimSun", 30),
                          bg=Color.b_c3).place(x=270, y=130)
            screen.update()
            time.sleep(1)
            login()
        cur.execute("CREATE TABLE IF NOT EXISTS info(id TEXT PRIMARY KEY, types TEXT, name TEXT,password TEXT)")
        list_info = cur.execute("SELECT * FROM info").fetchall()
        for i in list_info:
            # 判断用户是否之前注册过
            if i[0] == str(userid):
                connection.commit()
                cur.close()
                connection.close()
                tkinter.Label(screen, text="您已注册过!", font=("SimSun", 50), bg=Color.b_c3).place(x=500, y=130)
                screen.update()
                time.sleep(1)
                login()
            # 判断用户输入的用户名是否与其他用户的用户名重复
            if i[2] == username:
                connection.commit()
                cur.close()
                connection.close()
                tkinter.Label(screen, text="用户名已被注册!", font=("SimSun", 50), bg=Color.b_c3).place(x=500, y=130)
                screen.update()
                time.sleep(1)
                register()
        # 闯过上面三层识别之后就开始录入学生信息：
        cur.executemany('INSERT INTO info VALUES (?,?,?,?)', [[str(userid), user_types, username, password]])
        connection.commit()
        cur.close()
        connection.close()
        tkinter.Label(screen, text="注册成功!", font=("SimSun", 50), bg=Color.b_c3).place(x=500, y=130)
        screen.update()
        time.sleep(1)
        login()


def login(account=None, password=None, types2=None):
    """
    登陆界面，所有的参数在函数第一次调用时都为空，是无用的，在递归调用时才有用
    :param account: str 用户名
    :param password: str 密码
    :param types2: str 登陆错误的类型 "账号错误"/""密码错误"
    :return: None
    """
    tkinter.Label(screen, image=background_image).place(x=0, y=0)  # 设置背景
    # 用户输入账号于密码之前
    if account is None and password is None:
        # 登陆界面的背景图片

        # login_background_image = tkinter.PhotoImage(file='login_background_image.png')
        img = Image.open("data\login_background_image.png")  # 使用pillow打开PNG图片
        login_background_image = ImageTk.PhotoImage(img)
        # 创建标签并设置背景图
        label_login_background_image = tkinter.Label(screen, image=login_background_image)
        label_login_background_image.place(x=383, y=204)
        # 大标题
        tkinter.Label(screen, text="试卷题库管理系统", font=("SimSun", 45), bg='lightblue').place(x=430, y=100)
        # 先提醒用户注册
        tkinter.Label(screen, text="(如果没有账号，请先注册一个，或者使用:        登陆)", font=font1,
                      bg=Color.b_c3).place(x=400, y=440)
        tourist_mode_button = tkinter.Button(screen, text="游客模式", font=("SimSun", 12), command=login_as_tourist,
                                             activebackground=Color.sky_blue, bg=Color.white)  # 登陆按钮
        tourist_mode_button.place(x=810, y=440)
        register_button = tkinter.Button(screen, text="注册", font=font1, command=register,
                                         activebackground=Color.sky_blue, bg=Color.white)
        register_button.place(x=740, y=374)
        # 摆放提醒文字
        tkinter.Label(screen, text="账号", font=font1, bg=Color.b_c4).place(x=520, y=254)
        tkinter.Label(screen, text="密码", font=font1, bg=Color.b_c4).place(x=520, y=314)
        # 摆放输入框
        account_test = tkinter.Entry(screen, width=20, font=font1, highlightcolor=Color.b_c2,
                                     highlightthickness=1)  # 账号框
        account_test.place(x=600, y=254)
        password_test = tkinter.Entry(screen, width=20, font=font1, highlightcolor=Color.b_c2,
                                      highlightthickness=1)  # 密码框
        password_test.place(x=600, y=314)
        # 摆放登陆按钮，将用户输入的文本作为参数递归调用这个函数
        login_button = tkinter.Button(screen, text="登陆", font=font1,
                                      command=lambda: login(account=account_test.get(), password=password_test.get()),
                                      activebackground=Color.sky_blue,
                                      bg=Color.white)  # 登陆按钮
        login_button.place(x=570, y=374)
        # 如果密码有误或者账号有误时，提醒用户，然后在0.5s后让用户再次输入
        if types2 == "密码有误":
            tkinter.Label(screen, text="密码有误", font=("SimSun", 30), bg=Color.b_c3).place(x=623, y=375)
            screen.update()
            time.sleep(0.5)
            login()
        elif types2 == "账号有误":
            tkinter.Label(screen, text="账号有误", font=("SimSun", 30), bg=Color.b_c3).place(x=623, y=375)
            screen.update()
            time.sleep(0.5)
            login()
        screen.mainloop()  # 循环以显示窗口
    # 用户输入账号于密码之后
    else:
        connection = sqlite3.connect('data/usersinfo.db')  # 在同级目录下如果没有名为 usersinfo.db 的数据库，那么就创建一个
        cur = connection.cursor()  # 创建游标
        cur.execute("CREATE TABLE IF NOT EXISTS info(id TEXT PRIMARY KEY, types TEXT, name TEXT,password TEXT)")
        list_info = cur.execute("SELECT * FROM info").fetchall()
        for i in list_info:
            if i[2] == account:
                if i[3] == password:
                    global authority
                    authority = i[1]
                    main_menu(user_type=i[1])  # 登陆成功
                else:
                    login(types2="密码有误")
        login(types2="账号有误")


def login_as_tourist():
    """
    游客模式登陆
    :return: None
    """
    global authority
    authority = "student"  # 赋予游客与学生一样的权力
    main_menu(authority)


def main_menu(user_type):
    """
    函数注释:
    :param user_type: str 用户类型，必须是 student / teacher / admins / nt_authority ， 调试阶段，默认最高等级
    :return:None
    """

    # # 摆放背景图片
    # tkinter.Canvas(screen, height=714, width=1344, bg='white').place(x=0, y=0)  # 使用画布来遮住所有控件
    # tkinter.Label(screen, image=background_image).place(x=0, y=0)
    #
    # 摆放作者信息与版本号
    tkinter.Label(screen, text="欢迎使用Level7开发的题库管理系统。", font=font1,
                  foreground=Color.black).place(x=510, y=680)
    tkinter.Label(screen, text="Version: 1.0", font=font1, foreground=Color.black).place(x=1170, y=680)
    # 显示用户等级
    tkinter.Label(screen, text=f"权限: {user_type}.", font=font1,
                  foreground=Color.black).place(x=25, y=200)

    user_img = Image.open("data/dbt128.png")
    user_img = ImageTk.PhotoImage(user_img)
    user_label = tkinter.Label(screen, image=user_img)
    user_label.place(x=60, y=40)

    # 根据用户类型显示不同的按钮
    dict_authority = {"student": 1, "teacher": 2, "admins": 3, "nt_authority": 4}  # 权限表，一共4级
    # 所有人都能够查看题库
    tkinter.Button(screen, text="查看题库", font=font2, command=lambda: choose(find), bg=Color.grey,
                   activebackground="#ADD8E6", foreground=Color.black).place(x=50, y=260)
    # 新增关键词查找按钮
    # tkinter.Button(screen, text="关键词查找", font=font2, command=lambda: open_keyword_search(), bg=Color.b_c1,
    #                activebackground=Color.b_c2, foreground=Color.white).place(x=75, y=360)
    if dict_authority[user_type] >= 2:  # 老师
        tkinter.Button(screen, text="增加题目", font=font2, command=lambda: choose(add), bg=Color.grey,
                       activebackground="#ADD8E6", foreground=Color.black).place(x=50, y=320)
        tkinter.Button(screen, text="修改题目", font=font2, command=lambda: choose(sub_and_revise), bg=Color.grey,
                       activebackground="#ADD8E6", foreground=Color.black).place(x=50, y=380)
    if dict_authority[user_type] >= 3:  # 管理员
        tkinter.Button(screen, text="导入题目", font=font2, command=lambda: import_data(None), bg=Color.grey,
                       activebackground="#ADD8E6", foreground=Color.black).place(x=50, y=440)
        tkinter.Button(screen, text="导出题目", font=font2, command=lambda: derived_data(None), bg=Color.grey,
                       activebackground="#ADD8E6", foreground=Color.black).place(x=50, y=500)
    if dict_authority[user_type] >= 4:  # 超级管理员
        tkinter.Button(screen, text="导入用户信息", font=font2, command=lambda: import_member_information(None),
                       bg=Color.grey, activebackground="#ADD8E6", foreground=Color.black).place(x=20, y=560)

    # 公共的按钮
    # tkinter.Button(screen, text="结束", font=font2, command=sys.exit, bg=Color.b_c1, activebackground=Color.red,
    #                foreground=Color.white).place(x=1100, y=450)
    tkinter.Button(screen, text="退出登陆", font=font2, command=login, bg=Color.grey, activebackground="#ADD8E6",
                   foreground=Color.black).place(x=50, y=620)

    # 循环显示窗口
    screen.mainloop()


def import_data(file_path=None):
    """
    导入题目的界面，同时负责导入题目到数据库里面的功能
    :param file_path:str 在第一次调用时无用，在递归调用是文件路径
    :return: None
    """
    # 刷新窗口，返回上一级的按钮链接到主菜单
    global screen
    if file_path is None:
        screen = renovate()
        tkinter.Button(screen, text="←", font=("SimSun", 30), command=lambda: main_menu(authority), bg=Color.sky_blue,
                       activebackground=Color.sky_blue_out).place(x=5, y=5)
        # 提示字段
        tkinter.Label(screen, text="请输入要导入的文件名(带后缀),请确保在同一目录下!", font=font1, bg=Color.b_c3).place(
            x=400, y=200)
        # 输入框与确定按钮
        file_name_text = tkinter.Entry(screen, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        file_name_text.place(x=400, y=244)
        tkinter.Button(screen, text="确认导入", font=font1, command=lambda: import_data(file_name_text.get()),
                       bg=Color.b_c1, activebackground=Color.b_c2).place(x=400, y=280)
        screen.mainloop()
    # 递归调用时跑这一行，这时用户已经输入了文件名了
    if file_path:
        list_all_problem = []  # 存储所有读取到的格式化的问题的列表
        try:
            if file_path[-5:] == ".docx":
                # 从docx中读入题目并规整化
                doc = docx.Document(file_path)
                list_temp = []  # 临时列表，用来生成多维列表的小列表
                for paragraph in doc.paragraphs:
                    if paragraph.text != '':
                        list_temp.append(paragraph.text)
                    else:
                        list_all_problem.append(list_temp)
                        list_temp = []
            elif file_path[-4:] == ".txt":
                file1 = open(file_path, "r", encoding='utf-8').readlines()
                list_temp = []
                for i in file1:
                    i = i[:-1]
                    if i:
                        list_temp.append(i)
                    else:
                        list_all_problem.append(list_temp)
                        list_temp = []
            elif file_path[-5:] == ".xlsx":
                file1 = openpyxl.load_workbook(file_path).active
                for line in file1:  # 读取1-5行,如果是['A':'D']就是读取A-D列
                    list_temp = []
                    for val in line:
                        if val.value:
                            list_temp.append(val.value)
                    list_all_problem.append(list_temp)
            else:
                tkinter.Label(screen, text="文件格式有误!", font=font1, bg=Color.b_c3).place(x=500, y=150)
                screen.update()
                time.sleep(0.5)
                import_data()
        except docx.opc.exceptions.PackageNotFoundError:
            tkinter.Label(screen, text="找不到该文件!", font=font1, bg=Color.b_c3).place(x=500, y=150)
            screen.update()
            time.sleep(0.5)
            import_data()
        except FileNotFoundError:
            tkinter.Label(screen, text="找不到该文件!", font=font1, bg=Color.b_c3).place(x=500, y=150)
            screen.update()
            time.sleep(0.5)
            import_data()

        connection = sqlite3.connect('data/examination_questions.db')  # 在同级目录下如果没有名为 examination_questions.db 的数据库，那么就创建一个
        cur = connection.cursor()  # 创建游标
        # 创建一个表（CREATE TABLE）并判断是否存在同名的情况（IF NOT EXISTS）于文件test中
        cur.execute("CREATE TABLE IF NOT EXISTS radio(id INTEGER PRIMARY KEY, \
                            problem TEXT, options TEXT, answer TEXT)")  # 创建单选题表,包含id，问题，选项，答案四个数据类型
        cur.execute("CREATE TABLE IF NOT EXISTS multiple_choice(id INTEGER PRIMARY KEY, \
                            problem TEXT, options TEXT, answer TEXT)")  # 创建多选题表,包含id，问题，选项，答案四个数据类型
        cur.execute("CREATE TABLE IF NOT EXISTS true_or_false(id INTEGER PRIMARY KEY, \
                            problem TEXT, answer TEXT)")  # 创建判断题表,包含id，问题，答案三个数据类型
        cur.execute("CREATE TABLE IF NOT EXISTS fill_in_the_blank(id INTEGER PRIMARY KEY, \
                            problem TEXT, answer TEXT)")  # 创建填空题表,包含id，问题，答案三个数据类型
        cur.execute("CREATE TABLE IF NOT EXISTS subjective(id INTEGER PRIMARY KEY, \
                            problem TEXT, answer TEXT)")  # 创建主观题表,包含id，问题，答案三个数据类型
        # 创建各个题目的id，同种题目内部的id是唯一的，新的id是原来存在的最大的id+1
        id_radio = max([i[0] for i in cur.execute("SELECT id FROM radio").fetchall()] + [0]) + 1
        id_multiple_choice = max([i[0] for i in cur.execute("SELECT id FROM multiple_choice").fetchall()] + [0]) + 1
        id_true_or_false = max([i[0] for i in cur.execute("SELECT id FROM true_or_false").fetchall()] + [0]) + 1
        id_fill_in_the_blank = max([i[0] for i in cur.execute("SELECT id FROM fill_in_the_blank").fetchall()] + [0]) + 1
        id_subjective = max([i[0] for i in cur.execute("SELECT id FROM subjective").fetchall()] + [0]) + 1
        iter_temp = 0  # 用来生成题目的id
        # 接下来创建各种题目的列表，将所有问题分流成各种问题并存储到对应的数据库里面
        list_radio = []
        list_multiple_choice = []
        list_true_or_false = []
        list_fill_in_the_blank = []
        list_subjective = []
        # 维护上面几行代码的时候不要作死在没有deepcopy的情况下用连等初始化赋值哟~~
        for i in list_all_problem:
            # 这里的 len(i[0]) >= 5 本来是不用判断的，但是为了用户导入奇怪的文件时不报错还是判断了
            if len(i[0]) >= 5 and i[0][0:5] == "#单选题#":
                # 这里为了匹配enum的数据类型跟cur.executemany的子元素元组，还操作了一下
                list_radio.append((id_radio + iter_temp, i[0][5::], str(i[1:-1]), i[-1]))
            elif len(i[0]) >= 5 and i[0][0:5] == "#多选题#":
                list_multiple_choice.append((id_multiple_choice + iter_temp, i[0][5::], str(i[1:-1]), i[-1]))
            elif len(i[0]) >= 5 and i[0][0:5] == "#判断题#":
                i.insert(0, id_true_or_false + iter_temp)
                list_true_or_false.append((i[0], i[1][5::], i[-1]))
            elif len(i[0]) >= 5 and i[0][0:5] == "#填空题#":
                i.insert(0, id_fill_in_the_blank + iter_temp)
                list_fill_in_the_blank.append((i[0], i[1][5::], i[-1]))
            elif len(i[0]) >= 5 and i[0][0:5] == "#主观题#":
                i.insert(0, id_subjective + iter_temp)
                list_subjective.append((i[0], i[1][5::], i[-1]))
            iter_temp += 1  # 每读入一条题目id+1
        # 将规整化的数据存入数据库中
        cur.executemany('INSERT INTO radio VALUES (?,?,?,?)', list_radio)
        cur.executemany('INSERT INTO multiple_choice VALUES (?,?,?,?)', list_multiple_choice)
        cur.executemany('INSERT INTO true_or_false VALUES (?,?,?)', list_true_or_false)
        cur.executemany('INSERT INTO fill_in_the_blank VALUES (?,?,?)', list_fill_in_the_blank)
        cur.executemany('INSERT INTO subjective VALUES (?,?,?)', list_subjective)
        # 结束三连:确认修改,关闭游标,关闭链接
        connection.commit()
        cur.close()
        connection.close()
        tkinter.Label(screen, text="导入成功！", font=("SimSun", 50), bg=Color.b_c3).place(x=300, y=250)
        screen.update()
        time.sleep(1)
        main_menu(authority)


def derived_data(file_type=None, file_num=None):
    """
    导出题目界面，与其他函数几个一样，参数在递归调用时才有用
    :param file_type:要导出的文件类型
    :param file_num:要导出的题型
    :return:None,已经写入到文件里面了
    """
    screen = renovate()
    tkinter.Button(screen, text="←", font=("SimSun", 30), command=lambda: main_menu(authority), bg=Color.sky_blue,
                   activebackground=Color.sky_blue_out).place(x=5, y=5)
    # 初次进入该界面时：
    if file_type is None and file_num is None:
        value1 = tkinter.IntVar()  # 我们创建一个Int类型的容器存储单选框的值
        tkinter.Label(screen, text="↓请选择要导出的文件类型", font=font1, bg=Color.b_c3).place(x=300, y=200)
        tkinter.Radiobutton(screen, text='导出为docx', font=font1, variable=value1, value=0, bg=Color.b_c3).place(x=300,
                                                                                                                  y=250)
        tkinter.Radiobutton(screen, text='导出为txt', font=font1, variable=value1, value=1, bg=Color.b_c3).place(x=300,
                                                                                                                 y=300)
        tkinter.Radiobutton(screen, text='导出为xlsx', font=font1, variable=value1, value=2, bg=Color.b_c3).place(x=300,
                                                                                                                  y=350)
        tkinter.Label(screen, text="↓请选择要导出的题库", font=font1, bg=Color.b_c3).place(x=700, y=200)
        value2 = tkinter.IntVar()  # 我们创建一个Int类型的容器存储单选框的值
        tkinter.Radiobutton(screen, text='仅导出单选题', font=font1, variable=value2, value=0, bg=Color.b_c3).place(
            x=700, y=250)
        tkinter.Radiobutton(screen, text='仅导出多选题', font=font1, variable=value2, value=1, bg=Color.b_c3).place(
            x=700, y=300)
        tkinter.Radiobutton(screen, text='仅导出判断题', font=font1, variable=value2, value=2, bg=Color.b_c3).place(
            x=700, y=350)
        tkinter.Radiobutton(screen, text='仅导出填空题', font=font1, variable=value2, value=3, bg=Color.b_c3).place(
            x=700, y=400)
        tkinter.Radiobutton(screen, text='仅导出主观题', font=font1, variable=value2, value=4, bg=Color.b_c3).place(
            x=700, y=450)
        tkinter.Radiobutton(screen, text='导出所有题目题', font=font1, variable=value2, value=5, bg=Color.b_c3).place(
            x=700, y=500)
        tkinter.Button(screen, text="确认导出", font=font1, command=lambda: derived_data(value1.get(), value2.get()),
                       bg=Color.sky_blue, activebackground=Color.sky_blue_out).place(x=500, y=600)
    else:
        if file_num != 5:
            file = read_file(file_num)
        else:
            file = read_file(0) + read_file(1) + read_file(2) + read_file(3) + read_file(4)
        # 导出为docx
        if file_type == 0:
            doc = Document()
            for q in file:
                if len(q) > 3:
                    doc.add_paragraph(q[1])
                    for i in q[2]:
                        doc.add_paragraph(i)
                    doc.add_paragraph(q[-1] + "\n")
                else:
                    doc.add_paragraph(q[1])
                    doc.add_paragraph(q[-1] + "\n")
            doc.save('examination_questions.docx')
        # 导出为txt
        elif file_type == 1:
            with open("data/examination_questions.txt", 'a', encoding="utf-8") as f:
                for q in file:
                    if len(q) > 3:
                        f.write(q[1] + "\n")
                        for i in q[2]:
                            f.write(i + "\n")
                        f.write(q[-1] + "\n")
                        f.write("\n")
                    else:
                        f.write(q[1] + "\n")
                        f.write(q[-1] + "\n")
                        f.write("\n")
                f.close()
            f.close()
        # 导出为xlsx
        elif file_type == 2:
            wb = openpyxl.Workbook()
            ws = wb.active
            dict_temp = {0: "B", 1: "C", 2: "D", 3: "E", 4: "F", 5: "G", 7: "H"}
            for i in range(len(file)):
                ws[f'A{i + 1}'] = file[i][1]
                if len(file[i]) > 3:
                    for j in range(len(file[i][2])):
                        ws[f"{dict_temp[j]}{i + 1}"] = file[i][2][j]
                    ws[f"{dict_temp[len(file[i][2])]}{i + 1}"] = file[i][-1]
                else:
                    ws[f'B{i + 1}'] = file[i][-1]
            wb.save('examination_questions.xlsx')
        tkinter.Label(screen, text="导出成功", font=font2, bg=Color.b_c3).place(x=400, y=100)
        screen.update()
        time.sleep(0.5)
        main_menu(authority)


def choose(fun):
    """
    选择要进行操作的题库
    :param fun: 操作类型
    :return: None
    """
    # screen = renovate()
    # tkinter.Button(screen, text="←", font=("SimSun", 30), command=lambda: main_menu(authority), bg=Color.sky_blue,
    #                activebackground=Color.sky_blue_out).place(x=5, y=5)
    tkinter.Button(screen, text="单选题", font=font2, command=lambda: fun(0), bg=Color.b_c1,
                   activebackground=Color.b_c2, foreground=Color.white).place(x=350, y=50)
    tkinter.Button(screen, text="多选题", font=font2, command=lambda: fun(1), bg=Color.b_c1,
                   activebackground=Color.b_c2, foreground=Color.white).place(x=350, y=150)
    tkinter.Button(screen, text="判断题", font=font2, command=lambda: fun(2), bg=Color.b_c1,
                   activebackground=Color.b_c2, foreground=Color.white).place(x=350, y=250)
    tkinter.Button(screen, text="填空题", font=font2, command=lambda: fun(3), bg=Color.b_c1,
                   activebackground=Color.b_c2, foreground=Color.white).place(x=350, y=350)
    tkinter.Button(screen, text="主观题", font=font2, command=lambda: fun(4), bg=Color.b_c1,
                   activebackground=Color.b_c2, foreground=Color.white).place(x=350, y=450)


def add(file_num):
    """
    添加题目界面
    :param file_num: 要修改的文件的编号，变量dict_name
    :return: None
    """
    screen = renovate()
    # 上一级是添加题目的界面
    tkinter.Button(screen, text="←", font=("SimSun", 30), command=lambda: choose(add), bg=Color.sky_blue,
                   activebackground=Color.sky_blue_out).place(x=5, y=5)
    # 摆放各种控件↓↓↓↓
    tkinter.Label(screen, text="题干:", font=font1, bg=Color.b_c3).place(x=300, y=200)
    scr = tkinter.scrolledtext.ScrolledText(screen, font=font1, width=50, height=3, bg=Color.b_c3,
                                            highlightcolor=Color.b_c2, highlightthickness=1)
    scr.place(x=400, y=200)
    tkinter.Label(screen, text=f"题目id会自动分配", font=font1, bg=Color.b_c3).place(x=300, y=100)  # id
    # 将客观题与主观题分开操作
    if file_num < 2:
        tkinter.Label(screen, text="(请依次填写选项,多余的选项个数可不填)", font=font1, bg=Color.b_c3).place(x=300,
                                                                                                             y=140)
        # 选项
        tkinter.Label(screen, text="选项1:", font=font1, bg=Color.b_c3).place(x=300, y=304)
        tkinter.Label(screen, text="选项2:", font=font1, bg=Color.b_c3).place(x=300, y=344)
        tkinter.Label(screen, text="选项3:", font=font1, bg=Color.b_c3).place(x=300, y=384)
        tkinter.Label(screen, text="选项4:", font=font1, bg=Color.b_c3).place(x=300, y=424)
        tkinter.Label(screen, text="选项5:", font=font1, bg=Color.b_c3).place(x=300, y=464)
        tkinter.Label(screen, text="选项6:", font=font1, bg=Color.b_c3).place(x=300, y=504)
        choose_test1 = tkinter.Entry(screen, width=50, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        choose_test1.place(x=400, y=304)
        choose_test2 = tkinter.Entry(screen, width=50, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        choose_test2.place(x=400, y=344)
        choose_test3 = tkinter.Entry(screen, width=50, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        choose_test3.place(x=400, y=384)
        choose_test4 = tkinter.Entry(screen, width=50, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        choose_test4.place(x=400, y=424)
        choose_test5 = tkinter.Entry(screen, width=50, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        choose_test5.place(x=400, y=464)
        choose_test6 = tkinter.Entry(screen, width=50, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        choose_test6.place(x=400, y=504)
        # 选择题答案
        tkinter.Label(screen, text="答案:", font=font1, bg=Color.b_c3).place(x=300, y=544)
        account_test6 = tkinter.Entry(screen, width=50, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        account_test6.place(x=400, y=544)
        # 诡异的规整方式...
        tkinter.Button(screen, text="添加", font=font2, bg=Color.b_c1,
                       activebackground=Color.b_c2, command=lambda: rewrite_file(file_num, "add",
                                                                                 [scr.get(1.0, 'end'),
                                                                                  [i for i in [choose_test1.get(),
                                                                                               choose_test2.get(),
                                                                                               choose_test3.get(),
                                                                                               choose_test4.get(),
                                                                                               choose_test5.get(),
                                                                                               choose_test6.get()]
                                                                                   if i],
                                                                                  account_test6.get()])).place(
            x=500, y=600)
    else:
        # 主观题答案
        tkinter.Label(screen, text="答案:", font=font1, bg=Color.b_c3).place(x=300, y=400)
        scr2 = tkinter.scrolledtext.ScrolledText(screen, font=font1, width=50, height=3, bg=Color.b_c3,
                                                 highlightcolor=Color.b_c2, highlightthickness=1)
        scr2.place(x=400, y=400)
        tkinter.Button(screen, text="添加", font=font2, bg=Color.b_c1,
                       activebackground=Color.b_c2, command=lambda: rewrite_file(file_num, "add",
                                                                                 [scr.get(1.0, 'end'),
                                                                                  scr2.get(1.0, 'end')],
                                                                                 file_num)).place(x=500, y=600)


def sub_and_revise(file_num, iter1=0):
    """
    修改题目界面
    :param file_num: 操作的文件的序号
    :param iter1: 记录题1
    :return: None
    """
    screen = renovate()
    # 上一级是修改题目的界面
    tkinter.Button(screen, text="←", font=("SimSun", 30), command=lambda: choose(sub_and_revise), bg=Color.sky_blue,
                   activebackground=Color.sky_blue_out).place(x=5, y=5)
    list_q = read_file(file_num)  # [[id(int),题干(str),选项[list(str)],答案(str)],...]
    tkinter.Label(screen, text=f"题目id:{list_q[iter1][0]}(不可修改)", font=font1, bg=Color.b_c3).place(x=500,
                                                                                                        y=100)  # id
    tkinter.Label(screen, text="题干:", font=font1, bg=Color.b_c3).place(x=300, y=200)
    scr = tkinter.scrolledtext.ScrolledText(screen, font=font1, width=50, height=3, bg=Color.b_c3,
                                            highlightcolor=Color.b_c2, highlightthickness=1)
    scr.insert("end", list_q[iter1][1])  # 将题干插入到滚动文本框里面去1
    scr.place(x=400, y=200)
    if file_num < 2:  # 选择题
        tkinter.Label(screen, text="(可增加或者减少选项,多余的选项个数可不填)", font=font1, bg=Color.b_c3).place(x=300,
                                                                                                                 y=140)
        tkinter.Label(screen, text="选项1:", font=font1, bg=Color.b_c3).place(x=300, y=304)
        tkinter.Label(screen, text="选项2:", font=font1, bg=Color.b_c3).place(x=300, y=344)
        tkinter.Label(screen, text="选项3:", font=font1, bg=Color.b_c3).place(x=300, y=384)
        tkinter.Label(screen, text="选项4:", font=font1, bg=Color.b_c3).place(x=300, y=424)
        tkinter.Label(screen, text="选项5:", font=font1, bg=Color.b_c3).place(x=300, y=464)
        tkinter.Label(screen, text="选项6:", font=font1, bg=Color.b_c3).place(x=300, y=504)
        choose_test1 = tkinter.Entry(screen, width=50, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        choose_test1.insert(0, list_q[iter1][2][0])  # 放文本
        choose_test1.place(x=400, y=304)
        choose_test2 = tkinter.Entry(screen, width=50, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        choose_test2.insert(0, str(list_q[iter1][2][1]))
        choose_test2.place(x=400, y=344)
        choose_test3 = tkinter.Entry(screen, width=50, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        choose_test3.insert(0, str(list_q[iter1][2][2]))
        choose_test3.place(x=400, y=384)
        choose_test4 = tkinter.Entry(screen, width=50, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        if len(list_q[iter1][2]) > 3:
            choose_test4.insert(0, str(list_q[iter1][2][3]))
        choose_test4.place(x=400, y=424)
        choose_test5 = tkinter.Entry(screen, width=50, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        if len(list_q[iter1][2]) > 4:
            choose_test5.insert(0, str(list_q[iter1][2][4]))
        choose_test5.place(x=400, y=464)
        choose_test6 = tkinter.Entry(screen, width=50, font=font1, highlightcolor=Color.b_c2, highlightthickness=1)
        if len(list_q[iter1][2]) > 5:
            choose_test6.insert(0, str(list_q[iter1][2][5]))
        choose_test6.place(x=400, y=504)
        # 选择题答案
        tkinter.Label(screen, text="答案:", font=font1, bg=Color.b_c3).place(x=300, y=550)
        text5 = tkinter.Entry(screen, font=font1, width=30, highlightcolor=Color.b_c2, highlightthickness=1)  # 答案
        text5.insert(0, list_q[iter1][-1])
        text5.place(x=400, y=550)
    else:  # 主观题答案，放在原本选项的位置
        tkinter.Label(screen, text="答案:", font=font1, bg=Color.b_c3).place(x=300, y=400)
        scr2 = tkinter.scrolledtext.ScrolledText(screen, font=font1, width=50, height=3, bg=Color.b_c3,
                                                 highlightcolor=Color.b_c2, highlightthickness=1)
        scr2.insert("end", list_q[iter1][-1])  # 将题干插入到滚动文本框里面去1
        scr2.place(x=400, y=400)
    tkinter.Label(screen, text=f"第{iter1 + 1}题", font=font1, bg=Color.b_c3).place(x=400, y=100)
    # 左右查询按钮
    if iter1 < len(list_q) - 1:
        tkinter.Button(screen, text="下一题", font=("SimSun", 30), bg=Color.b_c1,
                       activebackground=Color.b_c2,
                       command=lambda: sub_and_revise(file_num, iter1 + 1)).place(x=800, y=600)  # 注意跳转的位置
    if iter1 > 0:
        tkinter.Button(screen, text="上一题", font=("SimSun", 30), bg=Color.b_c1,
                       activebackground=Color.b_c2,
                       command=lambda: sub_and_revise(file_num, iter1 - 1)).place(x=200, y=600)
    # 删除按钮
    tkinter.Button(screen, text="删除", font=("SimSun", 30), bg=Color.b_c1,
                   activebackground=Color.b_c2,
                   command=lambda: rewrite_file(file_num, "sub", id=list_q[iter1][0])).place(x=400, y=600)
    # 修改按钮
    if file_num < 2:
        tkinter.Button(screen, text="修改", font=("SimSun", 30), bg=Color.b_c1,
                       activebackground=Color.b_c2,
                       command=lambda: rewrite_file(file_num, "revise", [scr.get(1.0, 'end'),
                                                                         [i for i in [choose_test1.get(),
                                                                                      choose_test2.get(),
                                                                                      choose_test3.get(),
                                                                                      choose_test4.get(),
                                                                                      choose_test5.get(),
                                                                                      choose_test6.get()]
                                                                          if i], text5.get()],
                                                    id=list_q[iter1][0])).place(
            x=600, y=600)
    else:
        tkinter.Button(screen, text="修改", font=("SimSun", 30), bg=Color.b_c1,
                       activebackground=Color.b_c2,
                       command=lambda: rewrite_file(file_num, "revise",
                                                    list_content=[scr.get(1.0, 'end'), scr2.get(1.0, "end")],
                                                    id=list_q[iter1][0])).place(x=600, y=600)


def find(file_num, iter1=0):
    """
    查看题目界面
    :param file_num: 操作的文件的序号
    :param iter1: 记录题号
    :return: None
    """
    screen = renovate()
    # 上一级是查看题目的界面
    tkinter.Button(screen, text="←", font=("SimSun", 30), command=lambda: choose(find), bg=Color.sky_blue,
                   activebackground=Color.sky_blue_out).place(x=5, y=5)
    list_q = read_file(file_num)  # 读取对应的文件
    tkinter.Label(screen, text=f"第{iter1 + 1}题", font=font1, bg=Color.b_c3).place(x=400, y=100)
    tkinter.Label(screen, text=f"题目id:{list_q[iter1][0]}", font=font1, bg=Color.b_c3).place(x=500, y=100)  # id
    # 题干
    scr = tkinter.scrolledtext.ScrolledText(screen, font=font1, width=50, height=4, bg=Color.b_c3)
    scr.insert("end", list_q[iter1][1])  # 将题干插入到滚动文本框里面去1
    scr.configure(state='disabled')  # 设置为只读状态
    scr.place(x=300, y=150)
    if file_num < 2:  # 客观题，有选项，答案很短
        tkinter.Label(screen, text=f"答案{list_q[iter1][-1]}", font=font1, bg=Color.b_c3).place(x=300, y=550)
        option_num = len(list_q[iter1][2])
        if option_num < 5:  # 选项格式小于5个
            for i in range(len(list_q[iter1][2])):  # 选项
                if len(list_q[iter1][2][i]) <= 20:  # 短
                    tkinter.Label(screen, text=list_q[iter1][2][i], font=font1, width=50, bg=Color.b_c3).place(x=300,
                                                                                                               y=300 + i * 60)
                else:  # 长
                    scr = tkinter.scrolledtext.ScrolledText(screen, font=font1, width=50, height=1, bg=Color.b_c3)
                    scr.insert("end", list_q[iter1][2][i])  # 将题干插入到滚动文本框里面去
                    scr.configure(state='disabled')  # 设置为只读状态
                    scr.place(x=300, y=300 + i * 60)
        else:  # 选项个数超过5个，统一使用滚动文本框
            for i in range(len(list_q[iter1][2])):  # 选项
                scr = tkinter.scrolledtext.ScrolledText(screen, font=font1, width=23, height=1, bg=Color.b_c3)
                scr.insert("end", list_q[iter1][2][i])  # 将题干插入到滚动文本框里面去
                scr.configure(state='disabled')  # 设置为只读状态
                if i < 3:
                    scr.place(x=300, y=300 + i * 65)
                else:
                    scr.place(x=678, y=300 + (i - 3) * 65)
    else:  # 主观题，无选项，答案很长，我们把答案放在原本放选项的位置
        if len(list_q[iter1][-1]) <= 20:  # 可能使判断与填空的答案，还是很短
            tkinter.Label(screen, text=list_q[iter1][-1], font=font1, width=50, bg=Color.b_c3).place(x=300, y=350)
        else:
            scr2 = tkinter.scrolledtext.ScrolledText(screen, font=font1, width=50, height=4, bg=Color.b_c3)
            scr2.insert("end", list_q[iter1][-1])  # 将题干插入到滚动文本框里面去1
            scr2.configure(state='disabled')  # 设置为只读状态
            scr2.place(x=300, y=450)
    # 左右查询按钮
    if iter1 < len(list_q) - 1:
        tkinter.Button(screen, text="下一题", font=("SimSun", 30), bg=Color.b_c1,
                       activebackground=Color.b_c2,
                       command=lambda: find(file_num, iter1 + 1)).place(x=700, y=600)  # 注意跳转的位置
    if iter1 > 0:
        tkinter.Button(screen, text="上一题", font=("SimSun", 30), bg=Color.b_c1,
                       activebackground=Color.b_c2,
                       command=lambda: find(file_num, iter1 - 1)).place(x=400, y=600)


# 基础UI
screen = tkinter.Tk()  # 窗口初始化
screen.title("试卷题库管理系统")  # 设置标题
screen.geometry("1344x714+0+0")  # 设置窗口大小与初始化位置
screen.iconbitmap("data\python.ico")  # 设置窗口图标

'''
使用png设置图标
small_img = tkinter.PhotoImage(file="D:\Exam_Question_Bank_Management_System\data\dbt16.png")
large_img = tkinter.PhotoImage(file="D:\Exam_Question_Bank_Management_System\data\dbt32.png")
screen.iconphoto(False, large_img, small_img)
'''
screen.resizable(width=False, height=False)  # 禁止改变窗口大小!!!!!!

# background_image = tkinter.PhotoImage(file='background_image.png')  # 设置背景图片
# 使用pillow加载背景图片
img = Image.open("data/background_image.png")
background_image = ImageTk.PhotoImage(img)

#创建标签并设置背景图
background_label = tkinter.Label(screen, image=background_image)
background_label.place(x=0, y=0, relwidth=1, relheight=1)  # 设置标签占据整个窗口

font1 = ("SimSun", 16)  # 两种字体
font2 = ("SimSun", 25)
authority = None  # 全局变量，表示权限

if __name__ == '__main__':
    # 如果你是超级管理员，请直接运行下面这两行
    authority = "nt_authority"
    main_menu(user_type=authority)
    login()  # 开始运行！

"""
全局框架:               *import_member_information（导入用户信息）*
                                       ↓
                            login（登陆）/register（）注册
                                      ↓↓
            ↓-------------------------main_menu（主菜单）------------------------------------------------------↓
            |                           ↓↓                                                                   |
            |                    choose（选择操作的题型）                                                        |
            |                  ↓-------←→-------↓-------------------------------↓                            |
            ↓                  ↓                ↓                               ↓                            ↓
import_data（导入题目）     add（添加题目）     find（查看题目）         sub_and_revise（修改题目）        derived_data（导出题目）
                              |                 |                               |                       |
                              |                 |-----→ read_file（查看数据库）←---|←----------------------|
                              |                                                 |
                              |-----→------→rewrite_file（修改数据库）←----←-------|

加强:
1.支持图片
2.完全重复的题目会以不同的id被导入，显然这不是我们想要的情况
3.查询题目id跳转至对应位置
4.对与密码强度没有要求，也没有忘记密码的选项
5.游客模式的权限应该低于学生
6.连接服务器使用远程服务，数据库搭建在服务器上，避免学生直接删除数据库的情况
"""
